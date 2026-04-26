"""Generate the technical documentation .docx for the PFE.

Run:
    python scripts/generate_doc.py
Output:
    docs/Documentation_Technique_Arc_En_Ciel.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Documentation_Technique_Arc_En_Ciel.docx"

# ============================================================================
# Helpers
# ============================================================================


def set_cell_bg(cell, color_hex: str) -> None:
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = header_cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_i, row in enumerate(rows, start=1):
        for col_i, val in enumerate(row):
            cell = table.rows[row_i].cells[col_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


# ============================================================================
# Document content
# ============================================================================


def build_document() -> Document:
    doc = Document()

    # ----- Page de garde -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentation Technique\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Agent IA de recherche de subventions\npour l'association Arc En Ciel")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Projet de Fin d'Études\n")
    run.italic = True
    run.font.size = Pt(13)
    run = meta.add_run("Avril 2026")
    run.italic = True
    run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 1. Présentation du projet
    # ============================================================
    add_heading(doc, "1. Présentation du projet", 1)

    add_heading(doc, "1.1 Contexte", 2)
    add_para(doc,
        "Arc En Ciel est une association tunisienne (Sidi Daoued) fondée en 2012, "
        "œuvrant pour l'éducation, la formation et l'intégration des enfants à "
        "déficience mentale légère, spécialement les enfants atteints de trisomie 21. "
        "L'association construit actuellement un nouveau local et a besoin de financements "
        "pour son équipement (mobilier adapté, atelier pâtisserie, etc.).")

    add_heading(doc, "1.2 Objectif du projet", 2)
    add_para(doc,
        "Construire un agent IA backend qui automatise la recherche d'opportunités "
        "de subventions adaptées au profil de l'association, et envoie chaque semaine "
        "un récapitulatif des meilleures opportunités par email et par WhatsApp.")

    add_heading(doc, "1.3 Périmètre fonctionnel", 2)
    add_para(doc, "Le projet implémente l'Epic 1 du backlog :")
    add_bullet(doc, [
        "User Story 1 : récupération d'opportunités de subventions en ligne",
        "User Story 2 : vérification des critères d'éligibilité",
        "User Story 3 : classement des opportunités par pertinence",
        "User Story 4 : envoi d'instructions par email + WhatsApp",
    ])

    add_heading(doc, "1.4 Bénéficiaires", 2)
    add_table(doc,
        ["Acteur", "Rôle"],
        [
            ["Arc En Ciel (ONG)", "Bénéficiaire final — reçoit les opportunités hebdomadaires"],
            ["Enfants trisomiques", "Bénéficiaires ultimes — financements pour leur prise en charge"],
            ["Coordinateur de l'association", "Lit le digest et lance les candidatures"],
        ])

    doc.add_page_break()

    # ============================================================
    # 2. Périmètre métier
    # ============================================================
    add_heading(doc, "2. Périmètre métier", 1)

    add_heading(doc, "2.1 Mission de l'association", 2)
    add_para(doc,
        "« Association pour l'éducation, la formation et l'intégration des enfants "
        "à déficience mentale légère, spécialement les enfants trisomiques. »",
        italic=True)

    add_heading(doc, "2.2 Profil d'éligibilité visé", 2)
    add_table(doc,
        ["Critère", "Valeur"],
        [
            ["Localisation", "Sidi Daoued, Tunisie (MENA / Afrique du Nord)"],
            ["Type d'organisation", "ONG (Organisation Non Gouvernementale)"],
            ["Secteur principal", "Éducation, handicap, inclusion"],
            ["Année de fondation", "2012"],
            ["Bénéficiaires", "Enfants atteints de trisomie 21 et leurs familles"],
            ["Langues", "Arabe, Français"],
            ["Montant cible (min)", "5 000 USD"],
            ["Montant cible (max)", "50 000 USD"],
            ["Secteurs privilégiés", "Inclusion, handicap, éducation spécialisée, autonomie"],
            ["Portée géographique", "Locale, nationale, internationale"],
        ])

    add_heading(doc, "2.3 Projet ciblé pour la subvention", 2)
    add_para(doc,
        "Équipement et développement du nouveau local d'Arc En Ciel — deux volets :")
    add_bullet(doc, [
        "Achat de mobilier et équipements adaptés pour aménager les espaces "
        "d'accueil et d'apprentissage",
        "Mise en place d'un atelier pâtisserie complet (four, plan de travail, "
        "ustensiles, équipements de sécurité) pour développer l'autonomie et les "
        "compétences pratiques des enfants",
    ])

    add_heading(doc, "2.4 Règles métier appliquées par l'agent", 2)
    add_table(doc,
        ["Règle", "Implémentation"],
        [
            ["Maximum 5 opportunités / digest", "Limite top_n(limit=5) dans ranker.py"],
            ["Pertinence obligatoire", "Filtre NOT_ELIGIBLE → exclu du digest (status='rejected')"],
            ["Deadline non dépassée", "Filtre dans rank_all → status='expired' si deadline < aujourd'hui"],
            ["Langue prioritaire FR", "Mistral instruit de produire du français en priorité"],
            ["Hiérarchie de tri", "Complets > Incomplets ; Nouveaux > Déjà envoyés ; Score desc"],
            ["Anti-spam (dedup URL)", "Hash SHA256 de l'URL source en clé unique SQLite"],
        ])

    doc.add_page_break()

    # ============================================================
    # 3. Architecture
    # ============================================================
    add_heading(doc, "3. Architecture", 1)

    add_heading(doc, "3.1 Vue d'ensemble", 2)
    add_para(doc,
        "L'agent fonctionne en pipeline séquentiel à 4 étapes, déclenché par un "
        "cron hebdomadaire (GitHub Actions, lundi matin). Chaque étape consomme "
        "et enrichit la base SQLite locale.")

    add_heading(doc, "3.2 Diagramme de flux", 2)
    add_code(doc,
        "  ┌────────────────────────────────────────────────────────────────┐\n"
        "  │                  GitHub Actions (cron lundi 06:00 UTC)         │\n"
        "  └─────────────────────────────┬──────────────────────────────────┘\n"
        "                                │\n"
        "                                ▼\n"
        "  ┌────────────────────────────────────────────────────────────────┐\n"
        "  │                     python -m src.main                         │\n"
        "  └─────────────────────────────┬──────────────────────────────────┘\n"
        "                                │\n"
        "  ┌─────────────────────────────▼──────────────────────────────────┐\n"
        "  │ STEP 1 — Fetch (fetcher.py)                                    │\n"
        "  │   Tavily Search ──▶ DDG fallback ──▶ Pré-filtre URL/keywords   │\n"
        "  │   ──▶ Mistral extrait JSON ──▶ SQLite (dedup hash URL)         │\n"
        "  └─────────────────────────────┬──────────────────────────────────┘\n"
        "                                │\n"
        "  ┌─────────────────────────────▼──────────────────────────────────┐\n"
        "  │ STEP 2 — Eligibility (eligibility.py)                          │\n"
        "  │   Mistral juge ELIGIBLE / POTENTIALLY / NOT_ELIGIBLE           │\n"
        "  │   ──▶ status = 'evaluated' (ou 'rejected' si NOT_ELIGIBLE)     │\n"
        "  └─────────────────────────────┬──────────────────────────────────┘\n"
        "                                │\n"
        "  ┌─────────────────────────────▼──────────────────────────────────┐\n"
        "  │ STEP 3 — Ranking (ranker.py)                                   │\n"
        "  │   Filtre deadlines passées ──▶ Mistral score 0-100             │\n"
        "  │   ──▶ status = 'scored' (ou 'expired' si deadline < aujourd'hui)│\n"
        "  └─────────────────────────────┬──────────────────────────────────┘\n"
        "                                │\n"
        "  ┌─────────────────────────────▼──────────────────────────────────┐\n"
        "  │ STEP 4 — Digest (digest.py)                                    │\n"
        "  │   top_n bucket-sort 4 niveaux ──▶ Email Gmail SMTP             │\n"
        "  │                                ──▶ WhatsApp Twilio             │\n"
        "  │   ──▶ status = 'notified'                                      │\n"
        "  └────────────────────────────────────────────────────────────────┘"
    )

    add_heading(doc, "3.3 Composants principaux", 2)
    add_table(doc,
        ["Module", "Rôle"],
        [
            ["src/main.py", "Point d'entrée — orchestre les 4 étapes"],
            ["src/config.py", "Charge .env et profil JSON"],
            ["src/db.py", "Schéma SQLite, connexion, dédup"],
            ["src/search/", "Tavily principal + DuckDuckGo fallback"],
            ["src/agent/llm.py", "Builder partagé du LLM Mistral + throttling"],
            ["src/agent/prompts.py", "Prompts système + queries de recherche"],
            ["src/agent/fetcher.py", "User Story 1 — extraction d'opportunités"],
            ["src/agent/eligibility.py", "User Story 2 — vérification d'éligibilité"],
            ["src/agent/ranker.py", "User Story 3 — scoring + ranking"],
            ["src/notifications/", "Envoi email (SMTP) + WhatsApp (Twilio)"],
            ["src/templates/weekly_email.txt", "Template Jinja2 du digest"],
            ["config/association_profile.json", "Profil métier d'Arc En Ciel"],
            [".github/workflows/weekly.yml", "Cron GitHub Actions hebdo"],
        ])

    doc.add_page_break()

    # ============================================================
    # 4. Stack technique
    # ============================================================
    add_heading(doc, "4. Stack technique", 1)

    add_heading(doc, "4.1 Vue synthétique", 2)
    add_table(doc,
        ["Couche", "Technologie", "Version", "Rôle"],
        [
            ["Langage", "Python", "3.11+", "Runtime principal"],
            ["Framework Agent", "LangChain", "≥ 0.3.0", "Orchestration de l'agent IA"],
            ["LLM", "Mistral AI", "mistral-small-latest", "Génération + raisonnement"],
            ["Recherche web", "Tavily", "≥ 0.5.0", "Search optimisée pour agents IA"],
            ["Recherche fallback", "DuckDuckGo", "≥ 6.3.0", "Fallback sans clé API"],
            ["Stockage", "SQLite", "stdlib Python", "Persistance grants + dédup"],
            ["Templating", "Jinja2", "≥ 3.1.4", "Rendu du template email"],
            ["Email", "SMTP Gmail", "smtplib stdlib", "Envoi du digest"],
            ["WhatsApp", "Twilio (Sandbox)", "≥ 9.3.0", "Envoi WhatsApp"],
            ["Validation env", "python-dotenv", "≥ 1.0.1", "Chargement .env"],
            ["Cron / Hosting", "GitHub Actions", "—", "Exécution hebdomadaire gratuite"],
        ])

    add_heading(doc, "4.2 Choix du LLM — Mistral AI", 2)
    add_para(doc,
        "Trois LLM ont été évalués au cours du projet. Mistral a été retenu après "
        "élimination de Groq (limites de tokens trop strictes) et Gemini (pas "
        "d'accès au free tier depuis la Tunisie).")
    add_table(doc,
        ["Provider", "Limite gratuite", "Verdict"],
        [
            ["Groq (llama-3.3-70b)", "100k tokens/jour", "✗ Limite atteinte rapidement"],
            ["Groq (llama-3.1-8b)", "6k tokens/min", "✗ TPM trop bas"],
            ["Google Gemini 2.0", "1M TPM", "✗ Bloqué géographiquement (Tunisie hors free tier)"],
            ["Mistral AI (small)", "1 RPS / 500k TPM", "✓ Choix retenu — accessible, rapide, FR natif"],
        ])

    add_heading(doc, "4.3 Choix du moteur de recherche", 2)
    add_table(doc,
        ["Moteur", "Free tier", "Choix"],
        [
            ["Tavily", "1000 requêtes/mois", "✓ Principal — optimisé pour agents IA"],
            ["DuckDuckGo", "Illimité (sans clé)", "✓ Fallback en cas d'échec Tavily"],
            ["Brave Search", "2000 requêtes/mois", "✗ Non retenu (parsing manuel requis)"],
            ["SerpAPI", "100 requêtes/mois", "✗ Non retenu (trop limité)"],
        ])

    add_heading(doc, "4.4 Choix d'hébergement — GitHub Actions", 2)
    add_para(doc,
        "GitHub Actions est utilisé pour le scheduling hebdomadaire. Avantages :")
    add_bullet(doc, [
        "Gratuit pour les repos publics (illimité), 2000 min/mois pour les privés",
        "Cron natif via le déclencheur 'schedule'",
        "Stockage centralisé du code, des secrets et de l'historique d'exécution",
        "Possibilité de déclencher manuellement (workflow_dispatch)",
        "Logs d'exécution accessibles en ligne",
    ])

    doc.add_page_break()

    # ============================================================
    # 5. Schéma de base de données
    # ============================================================
    add_heading(doc, "5. Base de données SQLite", 1)

    add_heading(doc, "5.1 Schéma de la table grants", 2)
    add_table(doc,
        ["Colonne", "Type", "Contrainte", "Description"],
        [
            ["id", "INTEGER", "PK AUTOINCREMENT", "Identifiant unique"],
            ["url_hash", "TEXT", "NOT NULL UNIQUE", "Hash SHA-256 de l'URL (anti-doublons)"],
            ["title", "TEXT", "NOT NULL", "Titre de l'opportunité"],
            ["organization", "TEXT", "—", "Organisme financeur"],
            ["amount", "TEXT", "—", "Montant disponible (libre)"],
            ["deadline", "TEXT", "—", "Date limite (AAAA-MM-JJ)"],
            ["source_url", "TEXT", "NOT NULL", "URL d'origine"],
            ["description", "TEXT", "—", "Résumé en français"],
            ["language", "TEXT", "—", "fr | en | ar"],
            ["eligibility", "TEXT", "—", "Verdict + raison"],
            ["score", "REAL", "—", "Note 0-100"],
            ["score_reason", "TEXT", "—", "Explication du score"],
            ["status", "TEXT", "DEFAULT 'new'", "État dans le pipeline"],
            ["fetched_at", "TEXT", "NOT NULL", "Horodatage extraction"],
            ["notified_at", "TEXT", "—", "Horodatage envoi (NULL si jamais envoyé)"],
        ])

    add_heading(doc, "5.2 Index", 2)
    add_bullet(doc, [
        "idx_grants_status : sur status (filtrage rapide)",
        "idx_grants_score : sur score DESC (tri du top_n)",
        "idx_grants_deadline : sur deadline (filtrage des expirés)",
    ])

    add_heading(doc, "5.3 Workflow des statuts", 2)
    add_code(doc,
        "        ┌─────────┐\n"
        "        │   new   │  ← après fetch\n"
        "        └────┬────┘\n"
        "             │ eligibility\n"
        "             ▼\n"
        "      ┌────────────┐\n"
        "      │ evaluated  │\n"
        "      └─┬────────┬─┘\n"
        "        │        │ ranking\n"
        "        ▼        ▼\n"
        "  ┌──────────┐ ┌─────────┐ ┌─────────┐\n"
        "  │ rejected │ │ expired │ │ scored  │\n"
        "  └──────────┘ └─────────┘ └────┬────┘\n"
        "  (NOT_ELIGIBLE)  (deadline      │\n"
        "                   < today)      │ digest\n"
        "                                 ▼\n"
        "                           ┌──────────┐\n"
        "                           │ notified │\n"
        "                           └──────────┘"
    )

    add_heading(doc, "5.4 Statuts détaillés", 2)
    add_table(doc,
        ["Statut", "Quand ?", "Inclus dans top_n ?"],
        [
            ["new", "Après extraction LLM", "Non"],
            ["evaluated", "Après évaluation d'éligibilité (POTENTIALLY/ELIGIBLE)", "Non"],
            ["rejected", "Si Mistral retourne NOT_ELIGIBLE", "Non"],
            ["expired", "Si deadline < date du jour au moment du ranking", "Non"],
            ["scored", "Après scoring (score > 0)", "✓ Oui"],
            ["notified", "Après envoi dans un digest", "✓ Oui (catégorie 'old')"],
        ])

    doc.add_page_break()

    # ============================================================
    # 6. Algorithmes clés
    # ============================================================
    add_heading(doc, "6. Algorithmes clés", 1)

    add_heading(doc, "6.1 Pipeline d'extraction (fetcher.py)", 2)
    add_para(doc,
        "Pour chaque requête de la liste SEARCH_QUERIES (11 variantes FR/EN), "
        "le fetcher exécute :")
    add_bullet(doc, [
        "Search Tavily (5 résultats max) ; bascule sur DuckDuckGo si Tavily renvoie 0",
        "Pré-filtrage URL : exclusion de wikipedia.org, facebook.com, twitter.com, "
        "linkedin.com, youtube.com, /blog/, /news/, /about, /contact",
        "Pré-filtrage par mots-clés : exigence d'au moins un mot parmi {grant, "
        "subvention, appel à projet, funding, financement, fund, bourse, deadline, …}",
        "Throttle 5s entre appels LLM",
        "Appel Mistral pour extraire un JSON structuré (titre, organisme, montant, "
        "deadline, description, langue, éligibilité)",
        "Si extrait → upsert SQLite (le hash URL bloque les doublons)",
    ])

    add_heading(doc, "6.2 Évaluation d'éligibilité (eligibility.py)", 2)
    add_para(doc,
        "Pour chaque grant en status='new', Mistral reçoit le profil complet de "
        "l'association + la description du grant et retourne un verdict en JSON.")
    add_para(doc, "Règle d'or appliquée : « dans le doute, ne rejette pas »", italic=True)
    add_bullet(doc, [
        "ELIGIBLE : zone géo et secteur explicitement compatibles",
        "POTENTIALLY_ELIGIBLE : ambigu — l'utilisateur final décidera (cas par défaut)",
        "NOT_ELIGIBLE : exclusion explicite (ex: 'US-only', 'Réservé aux ONG françaises')",
    ])

    add_heading(doc, "6.3 Scoring et ranking (ranker.py)", 2)
    add_para(doc, "Pour chaque grant éligible, Mistral calcule un score 0-100 selon 4 critères :")
    add_table(doc,
        ["Critère", "Poids", "Détail du barème"],
        [
            ["Éligibilité", "40 %", "ELIGIBLE=40, POTENTIALLY=20, NOT_ELIGIBLE=0"],
            ["Alignement mission", "30 %", "Pertinence vs trisomie, inclusion, atelier pâtisserie, mobilier"],
            ["Montant", "15 %", "Idéal 5 000–50 000 USD = 15 ; raisonnable = 7 ; hors fourchette = 0"],
            ["Deadline", "15 %", "> 30 jours = 15 ; 7-30 = 10 ; < 7 = 5 ; passée = 0"],
        ])

    add_heading(doc, "6.4 Tri hiérarchique du digest (top_n)", 2)
    add_para(doc,
        "Le tri du top 5 obéit à 3 niveaux de priorité, du plus fort au plus faible :")
    add_table(doc,
        ["Niveau", "Critère", "Justification"],
        [
            ["1 (le + fort)", "Complet > Incomplet", "Les opportunités avec montant ET deadline connus sont plus actionnables"],
            ["2", "Nouveau ('scored') > Déjà envoyé ('notified')", "Privilégier la fraîcheur d'information"],
            ["3 (le + faible)", "Score décroissant", "Tiebreaker final"],
        ])

    add_para(doc, "Logique implémentée : 4 buckets, concatenation préservant l'ordre :")
    add_code(doc,
        "complete_new + complete_old + incomplete_new + incomplete_old\n"
        "└─ chaque bucket trié par score DESC ─┘\n"
        "résultat = liste[:5]"
    )

    add_heading(doc, "6.5 Throttling LLM", 2)
    add_para(doc,
        "Pour respecter le free tier de Mistral (1 requête/sec), un throttling "
        "manuel de 5 secondes entre chaque appel LLM est implémenté via une "
        "variable globale dans agent/llm.py (CALL_DELAY_SECONDS).")

    add_heading(doc, "6.6 Format du digest", 2)
    add_para(doc, "Email (template Jinja2) — extrait :")
    add_code(doc,
        'Objet : Arc En Ciel — Subventions de la semaine du {{ monday_date }}\n\n'
        'Bonjour,\n\n'
        '🥇 Opportunité 1 — {{ titre }}\n'
        '- Organisation : {{ organisme }}\n'
        '- Montant : {{ montant }}\n'
        '- Date limite : {{ deadline }}\n'
        '- Pourquoi ça correspond : {{ score_reason }}\n'
        '- Lien : {{ url }}\n'
        '\n'
        '🥈 Opportunité 2 — ... (idem)\n'
        '...\n\n'
        "L'agent Arc En Ciel 🌈"
    )
    add_para(doc, "WhatsApp — version compacte avec gras (*texte*) et italique (_texte_).")

    doc.add_page_break()

    # ============================================================
    # 7. Variables d'environnement
    # ============================================================
    add_heading(doc, "7. Variables d'environnement", 1)
    add_para(doc,
        "Toutes les valeurs sensibles (clés API, mots de passe) sont stockées dans "
        "un fichier .env local (jamais commité — protégé par .gitignore). En production "
        "(GitHub Actions), les valeurs proviennent des Secrets GitHub.")

    add_table(doc,
        ["Variable", "Source", "Description"],
        [
            ["MISTRAL_API_KEY", "console.mistral.ai", "Clé API Mistral (obligatoire)"],
            ["MISTRAL_MODEL", "fixe", "Modèle Mistral utilisé (mistral-small-latest)"],
            ["TAVILY_API_KEY", "app.tavily.com", "Clé API Tavily (obligatoire)"],
            ["SMTP_HOST", "fixe", "smtp.gmail.com"],
            ["SMTP_PORT", "fixe", "587"],
            ["SMTP_USER", "compte agent", "marzoukagent@gmail.com"],
            ["SMTP_APP_PASSWORD", "Google App Pwd", "16 caractères générés via myaccount.google.com/apppasswords"],
            ["EMAIL_FROM_NAME", "fixe", "« Agent Arc En Ciel »"],
            ["EMAIL_TO", "destinataire", "rahmaouerfelli555@gmail.com"],
            ["TWILIO_ACCOUNT_SID", "console.twilio.com", "Identifiant compte Twilio"],
            ["TWILIO_AUTH_TOKEN", "console.twilio.com", "Token d'authentification Twilio"],
            ["TWILIO_WHATSAPP_FROM", "fixe", "whatsapp:+14155238886 (sandbox)"],
            ["WHATSAPP_TO", "destinataire", "whatsapp:+216XXXXXXXX"],
            ["DB_PATH", "fixe", "data/grants.db"],
        ])

    doc.add_page_break()

    # ============================================================
    # 8. Dépendances Python
    # ============================================================
    add_heading(doc, "8. Dépendances Python (requirements.txt)", 1)
    add_table(doc,
        ["Paquet", "Version min", "Usage"],
        [
            ["langchain", "≥ 0.3.0", "Framework agent — abstractions LLM/messages"],
            ["langchain-mistralai", "≥ 0.2.0", "Connecteur LangChain pour Mistral AI"],
            ["langchain-community", "≥ 0.3.0", "Composants communautaires LangChain"],
            ["tavily-python", "≥ 0.5.0", "Client officiel Tavily"],
            ["duckduckgo-search", "≥ 6.3.0", "Search DuckDuckGo (fallback gratuit)"],
            ["pydantic", "≥ 2.9.0", "Validation de schémas (utilisé par LangChain)"],
            ["python-dotenv", "≥ 1.0.1", "Chargement des variables .env"],
            ["twilio", "≥ 9.3.0", "SDK officiel Twilio (WhatsApp)"],
            ["jinja2", "≥ 3.1.4", "Template engine pour l'email hebdo"],
            ["python-docx", "≥ 1.2.0", "Génération de cette documentation (script)"],
        ])

    doc.add_page_break()

    # ============================================================
    # 9. Déploiement et exécution
    # ============================================================
    add_heading(doc, "9. Déploiement et exécution", 1)

    add_heading(doc, "9.1 Modes d'exécution", 2)
    add_table(doc,
        ["Commande", "Effet"],
        [
            ["python -m src.main", "Pipeline complet : fetch → eligibility → rank → envoi"],
            ["python -m src.main --dry-run", "Pipeline complet sans envoi (test)"],
            ["python -m src.main --stats", "Affiche le contenu de la DB"],
            ["python -m src.main --reset-rejected", "Remet les NOT_ELIGIBLE en 'new' pour ré-éval"],
        ])

    add_heading(doc, "9.2 Workflow GitHub Actions", 2)
    add_para(doc, "Fichier : .github/workflows/weekly.yml — résumé :")
    add_code(doc,
        "name: Weekly grant digest\n"
        "on:\n"
        "  schedule:\n"
        "    - cron: '0 6 * * 1'    # tous les lundis à 06:00 UTC (07:00 Tunis)\n"
        "  workflow_dispatch:        # déclenchement manuel possible\n"
        "\n"
        "jobs:\n"
        "  run-agent:\n"
        "    runs-on: ubuntu-latest\n"
        "    timeout-minutes: 30\n"
        "    steps:\n"
        "      - actions/checkout@v4\n"
        "      - actions/setup-python@v5 (Python 3.11)\n"
        "      - pip install -r requirements.txt\n"
        "      - python -m src.main  (avec secrets injectés en env)\n"
        "      - upload-artifact (data/grants.db, 30 jours rétention)"
    )

    add_heading(doc, "9.3 Sécurité", 2)
    add_bullet(doc, [
        ".env protégé par .gitignore (jamais publié sur GitHub)",
        "Secrets GitHub stockés chiffrés dans Settings → Secrets and variables → Actions",
        "App Password Gmail : limité au scope 'mail', révocable à tout moment",
        "Twilio Sandbox : limité aux numéros qui ont rejoint via le code 'join xxxx'",
        "Aucun secret commité dans le code source",
        "DB SQLite locale uniquement — pas d'exposition publique",
    ])

    doc.add_page_break()

    # ============================================================
    # 10. Évolutions futures
    # ============================================================
    add_heading(doc, "10. Évolutions futures (hors périmètre PFE)", 1)
    add_para(doc,
        "Le périmètre actuel se limite à l'Epic 1 du backlog. Plusieurs améliorations "
        "ont été identifiées mais sont reportées :")
    add_bullet(doc, [
        "Migration de Twilio Sandbox vers Meta WhatsApp Cloud API (production)",
        "Interface web pour configurer les profils d'association (multi-tenant)",
        "Génération automatique des dossiers de candidature (templates)",
        "Suivi des candidatures déposées et de leur statut",
        "Tableau de bord analytique (taux d'éligibilité, taux de réussite, etc.)",
        "Alertes immédiates pour les opportunités très bien notées (sans attendre lundi)",
        "Élargissement du moteur de recherche à des bases spécialisées (Grants.gov, "
        "EU Funding & Tenders Portal, FDS, etc.)",
    ])

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"[OK] Document genere : {OUTPUT}")
    print(f"     Taille : {OUTPUT.stat().st_size // 1024} Ko")


if __name__ == "__main__":
    main()
